import os
import sys
import datetime
import uuid
import math

import dotenv
from groq import Groq
from pydub import AudioSegment
from docx import Document

# Models
model_whisper = "whisper-large-v3"
#model_gpt = "llama3-groq-70b-8192-tool-use-preview"
model_gpt = "llama-3.1-70b-versatile"


def split_audio(file_path):
    """
    Cutting the audio file so that it can be processed within the limits of the API.
    :param file_path: Path to audio file
    :return: Audio file cut into chunks
    """
    audio = AudioSegment.from_mp3(file_path)
    max_size = 25 * 1024 * 1024  # 25 MB
    chunk_length_ms = math.floor(0.9 * 1000 * (max_size / (audio.frame_rate * audio.frame_width)))
    chunks = [audio[i:i + chunk_length_ms] for i in range(0, len(audio), int(chunk_length_ms))]
    return chunks


def transcribe_audio(client,audio_chunks):
    """
    Convert audio files that have been cut into chunks into text.
    :param audio_chunks: Audio file cut into chunks
    :return: Audio file text
    """
    transcriptions = []
    for chunk in audio_chunks:
        temp_filename = "temp_audio_{}.mp3".format(uuid.uuid4())
        temp_audio_path = os.path.join("output", temp_filename)

        chunk.export(temp_audio_path, format="mp3", bitrate="192k")  # You can adjust the bitrate as needed
        file_size = os.path.getsize(temp_audio_path)

        if file_size > 25 * 1024 * 1024:
            os.remove(temp_audio_path)
            raise ValueError("Audio chunk is too large: {} bytes".format(file_size))

        with open(temp_audio_path, 'rb') as f:
            transcription = client.audio.transcriptions.create(file=f, model=model_whisper)
            
        if "请不吝点赞 订阅 转发 打赏支持明镜与点点栏目" in transcription.text:
            continue
        
        transcriptions.append(transcription.text)

        # Cleaning
        os.remove(temp_audio_path)

    return " ".join(transcriptions)


def abstract_summary_extraction(client,transcription):
    """
    From the audio file transcript, create a summary.
    :param transcription: Transcription of audio file
    :return: Abstract summary
    """
    response = client.chat.completions.create(
        model=model_gpt,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "您是一位训练有素、高度熟练的人工智能，能够理解和综合语言。根据以下文本，用一个抽象且简洁的段落进行总结。请尝试记住最重要的要点，提供一个连贯且可读的摘要，它可以帮助人们理解讨论的主要观点，而无需阅读全文。请避免不必要的细节或无关要点。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    response_dict = response.model_dump()
    return response_dict['choices'][0]['message']['content']


def key_points_extraction(client,transcription):
    """
    From the audio file transcript, create a list of key points.
    :param transcription: Transcription of audio file
    :return: Key points
    """
    response = client.chat.completions.create(
        model=model_gpt,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "您是一位训练有素的人工智能，专门从关键点中提取信息。根据以下文本，识别并列出已讨论或提及的主要要点。这些应该是对讨论本质至关重要的最重要的想法、结果或主题。您的目标是提供一份清单，以便某人可以快速阅读以了解已讨论的内容。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    response_dict = response.model_dump()
    return response_dict['choices'][0]['message']['content']


def action_item_extraction(client,transcription):
    """
    From the audio file transcript, create a list of action item.
    :param transcription: Transcription of audio file
    :return: Action item
    """
    response = client.chat.completions.create(
        model=model_gpt,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "您是一位训练有素的 AI，专门分析对话并提取需要执行的操作。根据以下文本，找出已达成一致或提及需要执行的任务、使命或操作。这些可能是分配给特定人员的任务，也可能是小组决定采取的常规操作。请列出这些操作的清晰简洁的清单。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    response_dict = response.model_dump()
    return response_dict['choices'][0]['message']['content']


def meeting_minutes(client,transcription):
    """
    Execution of all extractions.
    :param transcription: Transcription of audio file
    :return: List of all extractions
    """
    abstract_summary = abstract_summary_extraction(client,transcription)
    key_points = key_points_extraction(client,transcription)
    action_items = action_item_extraction(client,transcription)
    return {
        'complete_transcription': transcription,
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items
    }


def save_as_docx(minutes, filename, output_dir):
    """
    Save extraction as docx file.
    :param minutes: List of texts to be added
    :param filename: File name
    :param output_dir: Output directory
    """
    # Vérifie si le répertoire de sortie existe, sinon le crée
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def meeting_minutes_main(audio_file_path, choice, name_docx=None):
    """
    Main code for switching from an audio file to a transcription in a text file.
    :param audio_file_path: Path to the audio file (`.mp3`)
    :param choice: Choice between transcribing only or performing all actions ('Full' or 'Transcription')
    :param name_docx: Name of output text file (optional)
    """
    # Configuration
    dotenv.load_dotenv()
    api_key = os.getenv('GROQ_API_KEY')
    if api_key is None:
        print("API_KEY variable is not set: set it.")
        sys.exit(1)
    client = Groq(api_key= api_key,)

    # Split audio into chunks
    audio_chunks = split_audio(audio_file_path)

    # Always perform transcription
    transcription = transcribe_audio(client,audio_chunks)

    # Check user's choice
    if choice == 'Full':
        # If user chose 'full', perform all actions
        minutes = meeting_minutes(client,transcription)
        print("Minutes prepared.")
    elif choice == 'Transcription':
        # If user chose 'transcribe', only save the transcription
        minutes = {
            'complete_transcription': transcription
        }
        print("Transcription completed.")
    else:
        print("Invalid option. Exiting.")
        sys.exit(1)

    # Creating the output file
    output_dir = "output"
    if name_docx is None:
        now = datetime.datetime.now()
        formatted_date = now.strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"{output_dir}/meeting_minutes_{formatted_date}.docx"
    else:
        filename = f"{output_dir}/{name_docx}.docx"
    save_as_docx(minutes, filename, output_dir)
